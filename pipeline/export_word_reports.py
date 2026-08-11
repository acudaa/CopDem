#!/usr/bin/env python3
"""Export the two Austria analysis reports as Word documents
(reports/word/*.docx), for sharing outside a browser context.

Reuses the SAME stats-computation functions the HTML reports use
(analyze_obstacles_diff.py, analyze_lidar_raster_diff.py) rather than
recomputing anything independently — the numbers in the .docx files are
guaranteed to match the HTML reports exactly, not just "should match."

Charts are re-rendered as static matplotlib PNGs (not screenshots of the
interactive HTML/SVG charts) — appropriate for a static document, and it
means the by-type bar charts and histograms can use the exact same
underlying bin/stat data as the HTML version without fighting browser
screenshot cropping. Colors match the HTML reports' diverging blue/red
(#2a78d6 / #e34948) and the QGIS difference layer's convention, for
visual consistency across every deliverable in this project.

Requires python-docx and matplotlib (added to requirements.txt).
"""
import sys
from pathlib import Path

import matplotlib
matplotlib.use("Agg")  # no display needed, just PNG output
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt, RGBColor

sys.path.insert(0, str(Path(__file__).parent))
import analyze_obstacles_diff as obs_mod
import analyze_lidar_raster_diff as raster_mod

ROOT_DIR = Path(__file__).parent.parent
OUTPUT_DIR = ROOT_DIR / "reports" / "word"
TMP_CHART_DIR = OUTPUT_DIR / "_charts"

COLOR_NEG = "#2a78d6"  # blue - matches the HTML reports / QGIS diff layer
COLOR_POS = "#e34948"  # red
COLOR_SEQ = "#2a78d6"

HEADING_COLOR = RGBColor(0x0B, 0x0B, 0x0B)
MUTED_COLOR = RGBColor(0x60, 0x60, 0x60)


# ---------------------------------------------------------------------------
# Matplotlib chart rendering
# ---------------------------------------------------------------------------

def _style_axes(ax):
    for side in ("top", "right"):
        ax.spines[side].set_visible(False)
    ax.spines["left"].set_color("#c3c2b7")
    ax.spines["bottom"].set_color("#c3c2b7")
    ax.tick_params(colors="#52514e", labelsize=9)
    ax.grid(axis="y", color="#e1e0d9", linewidth=0.8, zorder=0)
    ax.set_axisbelow(True)


def plot_histogram(bins, out_path, xlabel="m", core_lo=None, core_hi=None,
                    core_start_idx=None, core_end_idx=None):
    """bins: list of (lo, hi, count).

    BUG FIXED HERE (found during PDF verification, not assumed away): bars
    are plotted by INDEX with equal screen width, never at true
    value-proportional width/position. The raster report's histogram()
    (analyze_lidar_raster_diff.py) inserts catch-all tail bins that can span
    a vastly wider value range than the core bins (e.g. a -362 to -7.6 tail
    bin next to ~1 m-wide core bins) - an earlier version of this function
    plotted bars at their true value width, which compressed the entire
    visible core distribution into a single spike next to two huge tail
    bars. This is the exact same problem already solved in the HTML report's
    svg_histogram() (see docs/lidar_comparison.md); this function now
    replicates that same index-based-bars-with-value-based-ticks approach
    instead of quietly doing something different in the Word version.

    core_lo/core_hi/core_start_idx/core_end_idx: when given (raster report),
    axis ticks are computed from these bounds, correctly interpolated onto
    equal-width bars via bin index - not by assuming bins[0]/bins[-1] are
    the chart's value extremes (they aren't once tail bins exist). When
    omitted (obstacle report, which uses simple min-to-max binning with no
    tails), the full bin range is used directly.
    """
    n_bins = len(bins)
    fig, ax = plt.subplots(figsize=(6.4, 2.6), dpi=150)

    if core_lo is None:
        core_lo, core_hi = bins[0][0], bins[-1][1]
        core_start_idx, core_end_idx = 0, n_bins - 1

    for i, (lo, hi, count) in enumerate(bins):
        mid = (lo + hi) / 2
        color = COLOR_NEG if mid < 0 else COLOR_POS
        is_tail = i < core_start_idx or i > core_end_idx
        ax.bar(i, count, width=0.92, align="edge", color=color,
               alpha=0.55 if is_tail else 1.0, zorder=3)

    def value_to_x(v):
        frac_of_core = (v - core_lo) / (core_hi - core_lo)
        return core_start_idx + frac_of_core * (core_end_idx - core_start_idx + 1)

    if core_lo < 0 < core_hi:
        ax.axvline(value_to_x(0), color="#c3c2b7", linewidth=1, zorder=2)

    # Snap whichever tick lands nearest zero to exactly 0 (same fix applied
    # to both HTML reports' svg_histogram()) - core_lo/core_hi are percentile
    # bounds, not symmetric about zero, so even spacing alone usually misses
    # an exact "0" label; readers need "no error" to be a visible tick, not
    # just the unlabeled zero-line drawn above.
    tick_fracs = (0, 0.25, 0.5, 0.75, 1)
    tick_vals = [core_lo + frac * (core_hi - core_lo) for frac in tick_fracs]
    if core_lo < 0 < core_hi:
        snap_i = min(range(len(tick_vals)), key=lambda i: abs(tick_vals[i]))
        tick_vals[snap_i] = 0
    tick_positions, tick_labels = [], []
    for val in tick_vals:
        tick_positions.append(value_to_x(val))
        tick_labels.append(f"{val:.0f}")
    ax.set_xticks(tick_positions)
    ax.set_xticklabels(tick_labels)
    ax.set_xlim(-0.2, n_bins + 0.2)

    _style_axes(ax)
    ax.set_xlabel(xlabel, fontsize=9, color="#52514e")
    ax.set_ylabel("cells", fontsize=9, color="#52514e")
    fig.tight_layout()
    fig.savefig(out_path, facecolor="white")
    plt.close(fig)


def plot_by_type_bar(type_stats, out_path, min_n_note=5):
    fig_h = max(1.8, 0.42 * len(type_stats) + 0.6)
    fig, ax = plt.subplots(figsize=(6.4, fig_h), dpi=150)
    labels = [s["type"] + (" †" if s["n"] < min_n_note else "") for s in type_stats]
    values = [s["mean"] for s in type_stats]
    colors = [COLOR_NEG if v < 0 else COLOR_POS for v in values]
    y_pos = range(len(type_stats))
    ax.barh(list(y_pos), values, color=colors, zorder=3, height=0.62)
    ax.set_yticks(list(y_pos))
    ax.set_yticklabels(labels, fontsize=8)
    ax.axvline(0, color="#c3c2b7", linewidth=1, zorder=2)
    _style_axes(ax)
    ax.grid(axis="x", color="#e1e0d9", linewidth=0.8, zorder=0)
    ax.set_xlabel("mean (m)", fontsize=9, color="#52514e")
    for i, v in enumerate(values):
        ax.text(v + (0.3 if v >= 0 else -0.3), i, f"{v:+.1f}", va="center",
                ha="left" if v >= 0 else "right", fontsize=7.5, color="#52514e")
    fig.tight_layout()
    fig.savefig(out_path, facecolor="white")
    plt.close(fig)


# ---------------------------------------------------------------------------
# docx helpers
# ---------------------------------------------------------------------------

def add_title(doc, text):
    p = doc.add_heading(text, level=0)
    return p


def add_intro(doc, text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.color.rgb = MUTED_COLOR
    return p


def add_table(doc, headers, rows, col_widths_cm=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for run in cells[i].paragraphs[0].runs:
                run.font.size = Pt(9)
    if col_widths_cm:
        # Two things are BOTH required for explicit widths to actually
        # render (verified by inspecting the generated XML, not assumed):
        # table.autofit=False sets <w:tblLayout type="fixed"/>, but the
        # column widths that a fixed-layout table actually renders from are
        # <w:tblGrid>'s <w:gridCol> elements - and setting row.cells[i].width
        # alone does NOT update those (confirmed: it left gridCol at its
        # default equal-split values, which is why the first version of this
        # function still rendered a squeezed, wrapped first column despite
        # setting cell widths). table.columns[i].width is the call that
        # correctly rewrites gridCol. Cell widths are set too, for direct
        # per-cell consistency, but gridCol is what actually controls layout.
        table.autofit = False
        for i, w in enumerate(col_widths_cm):
            table.columns[i].width = Cm(w)
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()  # spacer
    return table


def add_image(doc, path, width_in=6.2):
    doc.add_picture(str(path), width=Inches(width_in))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def fmt(v, decimals=1):
    return f"{v:,.{decimals}f}"


# ---------------------------------------------------------------------------
# Report 1: obstacle-based analysis (mirrors obstacles_dem_diff_analysis.html)
# ---------------------------------------------------------------------------

def build_obstacle_report():
    rows = obs_mod.load_rows()
    doc = Document()

    add_title(doc, "DEM vs. Obstacle Base — Analysis (Austria)")
    overall = {}
    for field_key in obs_mod.FIELDS:
        values = obs_mod.values_for(rows, field_key)
        overall[field_key] = obs_mod.compute_stats(values)

    n_total = len(rows)
    n_copdem = overall["error_5x5_median_m"]["n"]
    n_lidar = overall["error_5x5_median_lidar_m"]["n"]
    add_intro(
        doc,
        f"{n_total:,} obstacles (simple points, v1 scope) compared against two DEM sources — "
        f"Copernicus DEM GLO-30 ({n_copdem:,} with a valid 5×5-window sample) and "
        f"Austria's LiDAR-derived DHM ({n_lidar:,} with a valid 5×5-window sample). "
        f"Generated from obstacles_dem_diff.csv. Field definitions: docs/output_fields.md; "
        f"LiDAR vertical-datum source (GHA height, EPSG:5778): docs/lidar_comparison.md.",
    )

    doc.add_heading("Overview", level=1)
    headers = ["Field", "n", "mean", "median", "stdev", "p5", "p95", "min", "max"]
    overview_rows = []
    for field_key, meta in obs_mod.FIELDS.items():
        s = overall[field_key]
        overview_rows.append([
            f"{meta['group']} — {meta['label']}", f"{s['n']:,}", fmt(s["mean"]), fmt(s["median"]),
            fmt(s["stdev"]), fmt(s["p5"]), fmt(s["p95"]), fmt(s["min"]), fmt(s["max"]),
        ])
    add_table(doc, headers, overview_rows, col_widths_cm=[6, 1.25, 1.25, 1.25, 1.25, 1.25, 1.25, 1.25, 1.25])

    TMP_CHART_DIR.mkdir(parents=True, exist_ok=True)
    current_group = None
    for field_key, meta in obs_mod.FIELDS.items():
        if meta["group"] != current_group:
            current_group = meta["group"]
            doc.add_heading(current_group, level=1)

        s = overall[field_key]
        doc.add_heading(f"{meta['label']} ({meta['unit']})", level=2)
        doc.add_paragraph(meta["description"])

        headers = ["n", "mean", "median", "stdev", "p5–p95"]
        add_table(doc, headers, [[
            f"{s['n']:,}", fmt(s["mean"]), fmt(s["median"]), fmt(s["stdev"]),
            f"{fmt(s['p5'])} – {fmt(s['p95'])}",
        ]])

        values = obs_mod.values_for(rows, field_key)
        bins = obs_mod.histogram(values)
        hist_path = TMP_CHART_DIR / f"hist_{field_key}.png"
        plot_histogram(bins, hist_path, xlabel=f"{meta['label']} (m)")
        doc.add_paragraph("Distribution (whole dataset):").runs[0].font.bold = True
        add_image(doc, hist_path)

        type_stats = obs_mod.by_type_stats(rows, field_key)
        doc.add_paragraph("By obstacle type:").runs[0].font.bold = True
        type_headers = ["Obstacle type", "n", "mean", "median", "stdev", "min", "max"]
        type_rows = []
        for ts in type_stats:
            flag = " †" if ts["n"] < obs_mod.MIN_TYPE_N_FOR_NOTE else ""
            type_rows.append([
                ts["type"] + flag, f"{ts['n']:,}", fmt(ts["mean"]), fmt(ts["median"]),
                fmt(ts["stdev"]), fmt(ts["min"]), fmt(ts["max"]),
            ])
        add_table(doc, type_headers, type_rows, col_widths_cm=[5.5, 1.75, 1.75, 1.75, 1.75, 1.75, 1.75])

        bar_path = TMP_CHART_DIR / f"bar_{field_key}.png"
        plot_by_type_bar(type_stats, bar_path, obs_mod.MIN_TYPE_N_FOR_NOTE)
        add_image(doc, bar_path)

    doc.add_heading("Notes", level=1)
    doc.add_paragraph(
        f"† obstacle type with n < {obs_mod.MIN_TYPE_N_FOR_NOTE} — too few points for a "
        "stable mean. See docs/dem_extraction.md (Copernicus DEM) and docs/lidar_comparison.md "
        "(Austria LiDAR DEM) for sampling method and known limitations."
    )

    out_path = OUTPUT_DIR / "CopDem_Austria_Obstacle_Analysis.docx"
    doc.save(out_path)
    print(f"wrote {out_path}")
    return out_path


# ---------------------------------------------------------------------------
# Report 2: raster-based analysis (mirrors copdem_lidar_raster_diff_analysis.html)
# ---------------------------------------------------------------------------

def build_raster_report():
    data, native_shape, native_res = raster_mod.load_subsampled()
    import numpy as np
    valid = data[~np.isnan(data)]
    coverage_pct = 100 * valid.size / data.size
    stats = raster_mod.compute_stats(valid)
    bins, core_lo, core_hi, core_start_idx, core_end_idx = raster_mod.histogram(valid)
    native_px = native_shape[0] * native_shape[1]

    doc = Document()
    add_title(doc, "CopDEM vs. LiDAR — Raster Difference Analysis (Austria)")
    add_intro(
        doc,
        "Whole-country per-cell comparison, 10m grid (LiDAR's native resolution — CopDEM "
        "upsampled onto it), EGM2008. Not obstacle-based — see the companion obstacle-based "
        "report for point comparisons at obstacle locations. Generated from "
        "data/dem_diff_lidar_10m.tif. Raster build method and the LiDAR vertical-datum source "
        "(GHA height, EPSG:5778): docs/lidar_comparison.md.",
    )

    doc.add_heading("CopDEM − LiDAR (m)", level=1)
    doc.add_paragraph(
        f"Positive = CopDEM reads above LiDAR. Computed from a subsampled read (every "
        f"{raster_mod.SUBSAMPLE}th pixel in each direction — {stats['n']:,} sample cells "
        f"from {native_px:,} native 10m cells). {coverage_pct:.1f}% of the sampled area has "
        "valid data in both sources."
    )

    headers = ["mean", "median", "stdev", "p1", "p5", "p95", "p99", "min", "max"]
    add_table(doc, headers, [[
        fmt(stats["mean"]), fmt(stats["median"]), fmt(stats["stdev"]), fmt(stats["p1"]),
        fmt(stats["p5"]), fmt(stats["p95"]), fmt(stats["p99"]), fmt(stats["min"]), fmt(stats["max"]),
    ]])

    doc.add_paragraph(
        "Distribution — bins span the 1st–99th percentile at full resolution; the "
        "outermost bars (if present) are catch-alls for everything beyond, spanning the true "
        "min/max. No value clipping is applied to this raster — a handful of pixels in "
        "extreme terrain (e.g. steep peaks/ridgelines, where InSAR-derived DEMs like CopDEM are "
        "known to have large errors) can dominate the true min/max; see docs/lidar_comparison.md "
        "for a specific, individually-verified example near Grossglockner."
    ).runs[0].font.bold = True

    TMP_CHART_DIR.mkdir(parents=True, exist_ok=True)
    hist_path = TMP_CHART_DIR / "hist_raster.png"
    plot_histogram(bins, hist_path, xlabel="CopDEM − LiDAR (m)",
                    core_lo=core_lo, core_hi=core_hi,
                    core_start_idx=core_start_idx, core_end_idx=core_end_idx)
    add_image(doc, hist_path)

    out_path = OUTPUT_DIR / "CopDem_Austria_Raster_Analysis.docx"
    doc.save(out_path)
    print(f"wrote {out_path}")
    return out_path


if __name__ == "__main__":
    build_obstacle_report()
    build_raster_report()
